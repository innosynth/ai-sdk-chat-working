import os
import requests
import json
from pathlib import Path

# Test server URL
SERVER_URL = "http://localhost:5002"

def test_health():
    print("Testing health check...")
    response = requests.get(f"{SERVER_URL}/health")
    print(f"Health check response: {response.json()}")
    return response.status_code == 200

def test_text_file():
    print("\nTesting text file processing...")
    test_text = "This is a test text file.\nIt has multiple lines.\nAnd some special characters: !@#$%^&*()"
    
    # Create a temporary text file
    with open("test.txt", "w") as f:
        f.write(test_text)
    
    # Send the file to the server
    with open("test.txt", "rb") as f:
        files = {'file': f}
        data = {'fileType': 'txt', 'fileName': 'test.txt'}
        response = requests.post(f"{SERVER_URL}/process", files=files, data=data)
    
    # Clean up
    os.remove("test.txt")
    
    print(f"Text file processing response: {response.json()}")
    return response.status_code == 200

def test_pdf_file():
    print("\nTesting PDF file processing...")
    # Create a simple PDF file
    from reportlab.pdfgen import canvas
    
    c = canvas.Canvas("test.pdf")
    c.drawString(100, 750, "This is a test PDF file.")
    c.drawString(100, 700, "It contains multiple lines of text.")
    c.save()
    
    # Send the file to the server
    with open("test.pdf", "rb") as f:
        files = {'file': f}
        data = {'fileType': 'pdf', 'fileName': 'test.pdf'}
        response = requests.post(f"{SERVER_URL}/process", files=files, data=data)
    
    # Clean up
    os.remove("test.pdf")
    
    print(f"PDF file processing response: {response.json()}")
    return response.status_code == 200

def test_docx_file():
    print("\nTesting DOCX file processing...")
    from docx import Document
    
    # Create a simple DOCX file
    doc = Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.add_paragraph("It contains multiple paragraphs.")
    doc.save("test.docx")
    
    # Send the file to the server
    with open("test.docx", "rb") as f:
        files = {'file': f}
        data = {'fileType': 'docx', 'fileName': 'test.docx'}
        response = requests.post(f"{SERVER_URL}/process", files=files, data=data)
    
    # Clean up
    os.remove("test.docx")
    
    print(f"DOCX file processing response: {response.json()}")
    return response.status_code == 200

def test_csv_file():
    print("\nTesting CSV file processing...")
    import pandas as pd
    
    # Create a simple CSV file
    df = pd.DataFrame({
        'Name': ['Alice', 'Bob', 'Charlie'],
        'Age': [25, 30, 35],
        'City': ['New York', 'London', 'Paris']
    })
    df.to_csv("test.csv", index=False)
    
    # Send the file to the server
    with open("test.csv", "rb") as f:
        files = {'file': f}
        data = {'fileType': 'csv', 'fileName': 'test.csv'}
        response = requests.post(f"{SERVER_URL}/process", files=files, data=data)
    
    # Clean up
    os.remove("test.csv")
    
    print(f"CSV file processing response: {response.json()}")
    return response.status_code == 200

def main():
    print("Starting document processing tests...")
    
    # Test health check
    if not test_health():
        print("Health check failed!")
        return
    
    # Test different file types
    tests = [
        test_text_file,
        test_pdf_file,
        test_docx_file,
        test_csv_file
    ]
    
    results = []
    for test in tests:
        try:
            result = test()
            results.append(result)
        except Exception as e:
            print(f"Error in test: {str(e)}")
            results.append(False)
    
    # Print summary
    print("\nTest Summary:")
    for i, result in enumerate(results):
        test_name = tests[i].__name__
        status = "PASSED" if result else "FAILED"
        print(f"{test_name}: {status}")
    
    # Check if all tests passed
    if all(results):
        print("\nAll tests passed successfully!")
    else:
        print("\nSome tests failed. Please check the logs above.")

if __name__ == "__main__":
    main() 